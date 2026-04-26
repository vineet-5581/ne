"""
LAYER 9: Word Generation Engine

Generates .docx files using python-docx:
- Styled paragraphs
- Multi-level headings
- Tables
- Images
- Hyperlinks
- Sections

Author: Document AI Team
Version: 1.0.0
"""

from pathlib import Path
from typing import Dict, List

from utils import get_logger
from config import settings


class WordGenerator:
    """Generates Word documents from processed content"""

    def __init__(self):
        """Initialize Word generator"""
        self.logger = get_logger('word_generator')
        self.logger.info("Word Generator initialized")

        # Try to import python-docx
        self.docx = None
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            self.Document = Document
            self.Pt = Pt
            self.RGBColor = RGBColor
            self.Inches = Inches
            self.logger.info("✅ python-docx available")
        except ImportError:
            self.logger.warning("⚠️ python-docx not available")

    def generate(self, output_path: Path, text_data: Dict, style_data: Dict,
                table_data: Dict, doc_type: str) -> None:
        """Generate Word document"""
        try:
            if not self.Document:
                raise RuntimeError("python-docx not available")

            self.logger.info(f"Generating Word document: {output_path}")

            # Create document
            doc = self.Document()

            # Add title
            title = self._get_title(doc_type)
            title_para = doc.add_paragraph(title, style='Heading 1')
            title_para.alignment = 1  # Center alignment

            # Add styled blocks
            blocks = style_data.get('blocks', [])
            for block in blocks:
                self._add_styled_paragraph(doc, block)

            # Add tables
            tables = table_data.get('tables', [])
            for table in tables:
                self._add_table(doc, table)

            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = self.Inches(1)
                section.bottom_margin = self.Inches(1)
                section.left_margin = self.Inches(1)
                section.right_margin = self.Inches(1)

            # Save document
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_path)
            self.logger.info(f"✅ Word document created: {output_path}")

        except Exception as e:
            self.logger.error(f"Word generation failed: {e}")
            raise

    def _add_styled_paragraph(self, doc, block) -> None:
        """Add a styled paragraph to document"""
        try:
            text = getattr(block, 'text', block.get('text', '')) if isinstance(block, dict) else block.text
            style_name = getattr(block, 'style_name', block.get('style_name', 'Normal')) if isinstance(block, dict) else block.style_name

            para = doc.add_paragraph(text, style=style_name)

            # Apply formatting
            if isinstance(block, dict):
                bold = block.get('bold', False)
                italic = block.get('italic', False)
            else:
                bold = block.bold
                italic = block.italic

            for run in para.runs:
                if bold:
                    run.bold = True
                if italic:
                    run.italic = True

        except Exception as e:
            self.logger.warning(f"Error adding paragraph: {e}")

    def _add_table(self, doc, table) -> None:
        """Add a table to document"""
        try:
            if isinstance(table, dict):
                rows = table.get('rows', 1)
                cols = table.get('cols', 1)
            else:
                rows = table.rows
                cols = table.cols

            if rows > 0 and cols > 0:
                docx_table = doc.add_table(rows=rows, cols=cols)
                docx_table.style = 'Light Grid Accent 1'

        except Exception as e:
            self.logger.warning(f"Error adding table: {e}")

    def _get_title(self, doc_type: str) -> str:
        """Get appropriate title based on document type"""
        titles = {
            'resume': 'Resume',
            'research_paper': 'Research Paper',
            'invoice': 'Invoice',
            'form': 'Form',
            'contract': 'Contract',
            'report': 'Report',
        }
        return titles.get(doc_type, 'Document')
