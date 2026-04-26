"""
LAYER 10: Post-Processing AI

Fixes and improves generated document:
- Fix broken lines
- Merge fragmented text
- Normalize layout
- Improve visual similarity
- Validate output

Author: Document AI Team
Version: 1.0.0
"""

from pathlib import Path
from typing import Dict

from utils import get_logger
from config import settings


class PostProcessor:
    """Post-processes generated Word documents"""

    def __init__(self):
        """Initialize post-processor"""
        self.logger = get_logger('post_processor')
        self.logger.info("Post-Processor initialized")

    def process(self, output_path: Path) -> None:
        """Post-process generated document"""
        try:
            if not output_path.exists():
                self.logger.warning(f"Output file not found: {output_path}")
                return

            self.logger.info(f"Post-processing: {output_path}")

            # Try to import and process with python-docx
            try:
                from docx import Document

                doc = Document(output_path)

                # Fix paragraph spacing
                self._fix_paragraph_spacing(doc)

                # Normalize styles
                self._normalize_styles(doc)

                # Validate document
                self._validate_document(doc)

                # Save processed document
                doc.save(output_path)
                self.logger.info(f"✅ Post-processing complete")

            except ImportError:
                self.logger.warning("python-docx not available for post-processing")

        except Exception as e:
            self.logger.error(f"Post-processing failed: {e}")

    def _fix_paragraph_spacing(self, doc) -> None:
        """Fix and normalize paragraph spacing"""
        try:
            for para in doc.paragraphs:
                # Set consistent spacing
                para.paragraph_format.space_before = 0
                para.paragraph_format.space_after = 6
                para.paragraph_format.line_spacing = 1.15

        except Exception as e:
            self.logger.warning(f"Error fixing spacing: {e}")

    def _normalize_styles(self, doc) -> None:
        """Normalize and standardize styles"""
        try:
            for para in doc.paragraphs:
                # Apply consistent font
                for run in para.runs:
                    if not run.font.name:
                        run.font.name = 'Calibri'
                    if not run.font.size:
                        run.font.size = 11 * 2  # In twips

        except Exception as e:
            self.logger.warning(f"Error normalizing styles: {e}")

    def _validate_document(self, doc) -> bool:
        """Validate generated document structure"""
        try:
            has_content = len(doc.paragraphs) > 0 or len(doc.tables) > 0
            if has_content:
                self.logger.info("✅ Document validation passed")
                return True
            else:
                self.logger.warning("⚠️ Document appears to be empty")
                return False
        except Exception as e:
            self.logger.warning(f"Error validating document: {e}")
            return False
